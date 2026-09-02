"""Informe ejecutivo (2 hojas) de la Meta Comercial de septiembre 2026.

Destinatario: Coordinacion de Recaudo y Cartera.
Fuente de los datos: corrida del JOB_USP_Foto_Meta_Comercial_Mensual del
2026-09-01 07:40 sobre [Financiera].[Cartera_Meta_Comercial_Historico],
validada con analisis_meta_septiembre_2026.sql.

Estilo: Lineamientos_Visuales_y_Comunicacion_CUN_Word.md
"""
import sys

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout.reconfigure(encoding="utf-8", errors="replace")

SALIDA = "Informe_Ejecutivo_Meta_Comercial_Septiembre_2026.docx"

# ---- Paleta institucional CUN
AZUL_MARINO = RGBColor(0x0C, 0x23, 0x40)
AZUL_INST = RGBColor(0x1B, 0x36, 0x5D)
TURQUESA = RGBColor(0x00, 0x85, 0x9B)
VERDE_OSC = RGBColor(0x00, 0x7A, 0x33)
GRIS_INST = RGBColor(0x89, 0x8D, 0x8D)
TEXTO = RGBColor(0x22, 0x22, 0x22)
BLANCO = RGBColor(0xFF, 0xFF, 0xFF)

HEX_MARINO = "0C2340"
HEX_MANZANA = "84BD00"
HEX_ZEBRA = "F8F9FA"
HEX_CALLOUT = "F0F4F8"

TIT = "Montserrat"
CUERPO = "Open Sans"

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = CUERPO
normal.font.size = Pt(10)
normal.font.color.rgb = TEXTO
normal._element.rPr.rFonts.set(qn("w:eastAsia"), CUERPO)

for s in doc.sections:
    s.top_margin = Cm(1.6)
    s.bottom_margin = Cm(1.3)
    s.left_margin = Cm(2.2)
    s.right_margin = Cm(2.0)


# ------------------------------------------------------------------ helpers
def _fuente(run, nombre, size, bold=False, italic=False, color=None):
    run.font.name = nombre
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn("w:eastAsia"), nombre)
    return run


def shade(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexc)
    tcPr.append(shd)


def borde_izq(par, hexc, ancho=24):
    """Barra vertical de acento a la izquierda del parrafo (callout)."""
    pPr = par._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(ancho))
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), hexc)
    pbdr.append(left)
    pPr.append(pbdr)


def fondo_par(par, hexc):
    pPr = par._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexc)
    pPr.append(shd)


def linea_acento(hexc=HEX_MANZANA, ancho=24, after=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(after)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), str(ancho))
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), hexc)
    pbdr.append(bot)
    pPr.append(pbdr)
    return p


def par(texto="", size=10, bold=False, italic=False, color=TEXTO,
        fuente=CUERPO, after=5, before=0, align=None, interlineado=1.1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.line_spacing = interlineado
    if align is not None:
        p.alignment = align
    if texto:
        _fuente(p.add_run(texto), fuente, size, bold, italic, color)
    return p


def h1(texto):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(3)
    _fuente(p.add_run(texto), TIT, 12.5, bold=True, color=AZUL_MARINO)
    return p


def vineta(texto, negrita_hasta=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.55)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.08
    _fuente(p.add_run("▪  "), CUERPO, 10, bold=True, color=VERDE_OSC)
    if negrita_hasta:
        _fuente(p.add_run(negrita_hasta), CUERPO, 9.5, bold=True, color=AZUL_INST)
    _fuente(p.add_run(texto), CUERPO, 9.5, color=TEXTO)
    return p


def cell_text(cell, texto, bold=False, color=TEXTO, size=9, align=None, fuente=CUERPO):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    if align is not None:
        p.alignment = align
    _fuente(p.add_run(str(texto)), fuente, size, bold, False, color)


def tabla(headers, filas, anchos=None, size=8.5, alinear_der=None):
    """Tabla institucional: encabezado azul marino, zebra suave, sin bordes pesados."""
    alinear_der = alinear_der or []
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, hh in enumerate(headers):
        c = t.rows[0].cells[j]
        cell_text(c, hh, bold=True, color=BLANCO, size=size,
                  fuente=TIT,
                  align=WD_ALIGN_PARAGRAPH.RIGHT if j in alinear_der else None)
        shade(c, HEX_MARINO)
    for i, fila in enumerate(filas):
        cells = t.add_row().cells
        for j, val in enumerate(fila):
            neg = j == 0 and str(val).startswith("Total")
            cell_text(cells[j], val, bold=neg, size=size,
                      align=WD_ALIGN_PARAGRAPH.RIGHT if j in alinear_der else None)
            if i % 2 == 1:
                shade(cells[j], HEX_ZEBRA)
    if anchos:
        for row in t.rows:
            for j, w in enumerate(anchos):
                row.cells[j].width = Cm(w)
    return t


def callout(titulo, lineas, hexc_borde=HEX_MANZANA):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent = Cm(0.15)
    borde_izq(p, hexc_borde)
    fondo_par(p, HEX_CALLOUT)
    _fuente(p.add_run(titulo), TIT, 10.5, bold=True, color=AZUL_MARINO)
    for k, ln in enumerate(lineas):
        q = doc.add_paragraph()
        q.paragraph_format.space_before = Pt(0)
        q.paragraph_format.space_after = Pt(6 if k == len(lineas) - 1 else 1)
        q.paragraph_format.left_indent = Cm(0.15)
        q.paragraph_format.line_spacing = 1.08
        borde_izq(q, hexc_borde)
        fondo_par(q, HEX_CALLOUT)
        _fuente(q.add_run(ln), CUERPO, 9.5, color=TEXTO)


def kpi_row(items):
    """Fila de metricas destacadas: cifra grande turquesa + etiqueta gris."""
    t = doc.add_table(rows=2, cols=len(items))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, (cifra, etiqueta) in enumerate(items):
        cell_text(t.rows[0].cells[j], cifra, bold=True, color=TURQUESA, size=15,
                  align=WD_ALIGN_PARAGRAPH.CENTER, fuente=TIT)
        cell_text(t.rows[1].cells[j], etiqueta, color=GRIS_INST, size=8,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
    return t


def nota(texto):
    par(texto, size=8.5, italic=True, color=GRIS_INST, after=3)


# ------------------------------------------------------------------ encabezado
enc = doc.sections[0].header
enc.is_linked_to_previous = False
pe = enc.paragraphs[0]
pe.alignment = WD_ALIGN_PARAGRAPH.RIGHT
_fuente(pe.add_run("Corporación Unificada Nacional de Educación Superior - CUN"),
        CUERPO, 8, color=GRIS_INST)

pie = doc.sections[0].footer
pie.is_linked_to_previous = False
pp = pie.paragraphs[0]
pp.alignment = WD_ALIGN_PARAGRAPH.LEFT
_fuente(pp.add_run("Vicerrectoría de Servicios Digitales · Analítica financiera   "
                   "|   VIGILADA MINEDUCACIÓN"), CUERPO, 8, color=GRIS_INST)

# ------------------------------------------------------------------ HOJA 1
par("INFORME EJECUTIVO", size=9.5, bold=True, color=TURQUESA, fuente=TIT, after=1)
par("Meta Comercial de Cartera — Septiembre 2026", size=19, bold=True,
    color=AZUL_MARINO, fuente=TIT, after=2)
linea_acento(after=6)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
_fuente(p.add_run("Dirigido a: "), CUERPO, 9, bold=True, color=AZUL_INST)
_fuente(p.add_run("Óscar Penagos — Coordinación de Recaudo y Cartera     "),
        CUERPO, 9, color=TEXTO)
_fuente(p.add_run("Corte: "), CUERPO, 9, bold=True, color=AZUL_INST)
_fuente(p.add_run("1 de septiembre de 2026     "), CUERPO, 9, color=TEXTO)
_fuente(p.add_run("Elaboró: "), CUERPO, 9, bold=True, color=AZUL_INST)
_fuente(p.add_run("Analítica financiera CUN"), CUERPO, 9, color=TEXTO)

callout(
    "Conclusión",
    ["La meta de septiembre quedó conformada por 56.269 obligaciones de 24.474 estudiantes, "
     "por $16.527,3 millones. La generación automática se ejecutó sin novedades y los "
     "controles de integridad cuadran en su totalidad. Tres cuartas partes del saldo (74,3%) "
     "corresponde a cartera que ya venía en gestión y el 49% de las obligaciones lleva los "
     "cuatro meses del ciclo sin salir: ese es el núcleo del esfuerzo de recuperación."],
)

h1("1.  Resultado de la generación automática")
par("Ejecutamos el proceso programado el 1 de septiembre a las 07:40, con una duración de "
    "8 minutos 58 segundos y finalización correcta. Antes de refrescar la meta se tomó la foto "
    "mensual de respaldo de 65.456 registros, de modo que el estado con el que cerró agosto "
    "queda preservado para comparaciones posteriores.", after=4)
vineta("15.405 obligaciones nuevas más 40.864 acumuladas equivalen exactamente a las 56.269 "
       "marcadas en septiembre; el histórico no presenta créditos duplicados sobre 80.861 "
       "registros, ni marcas de mes ausentes o repetidas.", "Controles de integridad: ")

h1("2.  Tamaño de la meta")
kpi_row([("56.269", "OBLIGACIONES"),
         ("24.474", "ESTUDIANTES"),
         ("$16.527", "MILLONES"),
         ("$293.720", "TICKET PROMEDIO")])
par("Cada estudiante concentra en promedio 2,30 cuotas en mora: un solo contacto efectivo "
    "resuelve más de dos registros de la meta.", before=5, after=4)

h1("3.  Composición y movimiento del mes")
tabla(
    ["Concepto", "Obligaciones", "Estudiantes", "Saldo ($ MM)", "% saldo"],
    [["Cartera de arrastre (ya venía en la meta)", "40.864", "15.007", "12.274,6", "74,3%"],
     ["Ingresos nuevos de septiembre", "15.405", "15.324", "4.252,7", "25,7%"],
     ["Total meta septiembre", "56.269", "24.474", "16.527,3", "100,0%"]],
    anchos=[7.4, 2.5, 2.4, 2.5, 1.6],
    alinear_der=[1, 2, 3, 4],
)
par("", after=2)
vineta("9.289 obligaciones de 6.248 estudiantes, por $2.769,5 millones, salieron de la meta "
       "entre agosto y septiembre por pago o normalización.", "Salidas del mes: ")
vineta("27.619 obligaciones ($8.808,3 millones, el 49,1% del total) permanecen en la meta "
       "desde junio sin haber salido ningún mes.", "Permanencia crítica: ")
vineta("al 2 de septiembre ya salieron 1.793 obligaciones de 1.473 estudiantes "
       "por $479,7 millones.", "Avance temprano: ")

par("Periodos académicos con mayor participación", size=10, bold=True, color=AZUL_INST,
    fuente=TIT, after=3, before=8)
tabla(
    ["Periodo", "26V02", "26V01", "26V03", "26V04", "26ES3", "2026C", "Resto"],
    [["Obligaciones", "11.792", "11.474", "9.145", "6.307", "3.343", "2.918", "11.290"],
     ["Saldo ($ MM)", "3.052,2", "3.135,6", "1.893,5", "1.470,6", "982,1", "1.042,1", "4.951,2"]],
    anchos=[2.8, 1.9, 1.9, 1.9, 1.9, 1.9, 1.9, 1.9],
    alinear_der=[1, 2, 3, 4, 5, 6, 7],
)

# ------------------------------------------------------------------ HOJA 2
doc.add_page_break()

h1("4.  Priorización de la gestión")
par("La meta se reparte en cuatro cuartiles por saldo acumulado y antigüedad de vencimiento: "
    "Q1 agrupa lo más vencido y Q4 lo más reciente. Cada cuartil concentra exactamente el 25% "
    "del saldo, de modo que las cuatro cargas de trabajo son equivalentes.", after=4)

tabla(
    ["Cuartil", "Obligaciones", "Saldo ($ MM)", "% saldo", "Lectura de gestión"],
    [["Q1", "12.471", "4.131,4", "25,0%", "Mayor antigüedad — prioridad de contacto"],
     ["Q2", "13.653", "4.132,1", "25,0%", "Mora consolidada"],
     ["Q3", "15.231", "4.131,9", "25,0%", "Mora intermedia"],
     ["Q4", "14.914", "4.131,8", "25,0%", "Vencimiento reciente — mayor recuperabilidad"]],
    anchos=[1.7, 2.4, 2.4, 1.6, 8.3],
    alinear_der=[1, 2, 3],
)

par("Clasificación académica del deudor", size=10, bold=True, color=AZUL_INST,
    fuente=TIT, after=3, before=7)
tabla(
    ["Marca académica", "Obligaciones", "Estudiantes", "Saldo ($ MM)", "% saldo"],
    [["Periodo en curso", "24.955", "15.125", "6.529,0", "39,5%"],
     ["Sin registro de clase", "13.059", "3.497", "3.696,6", "22,4%"],
     ["Gestionable", "10.195", "3.396", "3.533,0", "21,4%"],
     ["Periodo perdido, prioridad alta", "7.949", "2.432", "2.729,5", "16,5%"],
     ["Periodo no ha iniciado", "111", "108", "39,1", "0,2%"]],
    anchos=[7.4, 2.5, 2.4, 2.5, 1.6],
    alinear_der=[1, 2, 3, 4],
)
nota("7.535 obligaciones por $2.511,2 millones suman perfil crediticio adverso y deben "
     "escalarse dentro de su marca.")

h1("5.  Antigüedad de la mora y concentración")
tabla(
    ["Rango de mora", "1–30", "31–60", "61–90", "91–120", "121–150", "151–360"],
    [["Obligaciones", "15.432", "8.811", "9.531", "8.695", "6.949", "6.846"],
     ["% del saldo", "25,8%", "13,6%", "16,8%", "16,3%", "13,3%", "14,2%"]],
    anchos=[3.4, 2.2, 2.2, 2.2, 2.2, 2.2, 2.2],
    alinear_der=[1, 2, 3, 4, 5, 6],
)
vineta("solo 5 obligaciones superan los 360 días. La meta es cartera joven y por tanto "
       "recuperable; el esfuerzo debe concentrarse antes de que los tramos de 121 a 360 "
       "días sigan creciendo.", "Sin cola larga: ")
vineta("el 99,4% de los deudores figura como estudiante activo: es población vigente, no "
       "desertores. La contactabilidad es completa en correo, con apenas 9 registros sin "
       "celular y 24 sin WhatsApp sobre 56.269.", "Población y contacto: ")

h1("6.  Observaciones de la Analítica para la Coordinación")
vineta("30,4% de las obligaciones (5.009 estudiantes, $3.696,6 millones) registra promedio "
       "acumulado inferior a 1,55 y queda como «sin registro de clase». Son matriculados sin "
       "calificación; sugerimos validarlo con Registro y Control antes de definir el guion de "
       "cobro de ese grupo.", "Calidad del dato académico: ")
vineta("el 85,7% de las obligaciones tiene promedio acumulado; el resto se prioriza por "
       "calendario y conexión a plataforma.", "Cobertura académica: ")
vineta("la foto de respaldo de agosto se tomó el día 27 y no el día 1, por lo que ese primer "
       "punto no es comparable con el de septiembre; desde este mes la serie queda homogénea.",
       "Serie histórica: ")

par("", after=2)
callout(
    "Recomendación",
    ["Priorizar el cruce entre Q1 y la marca «periodo perdido, prioridad alta» con perfil "
     "crediticio adverso: es el segmento donde coinciden mayor antigüedad, deterioro "
     "académico y deterioro crediticio, abordando la gestión por estudiante y no por "
     "obligación."],
    hexc_borde="00859B",
)

par("", after=2)
nota("Fuente: [Financiera].[Cartera_Meta_Comercial_Historico], extracción del 2 de "
     "septiembre de 2026. Cifras en millones de pesos colombianos.")

doc.save(SALIDA)
print(f"OK -> {SALIDA}")
