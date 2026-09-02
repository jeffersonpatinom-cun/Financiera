"""Identidad visual CUN para informes ejecutivos en Word.

Implementa Lineamientos_Visuales_y_Comunicacion_CUN_Word.md sobre python-docx.
Lo usan generar_informe_meta_septiembre.py y generar_informe_agosto.py; cualquier
informe nuevo debe importarlo en vez de recopiar los helpers.

Uso tipico:

    import sys
    sys.path.insert(0, ".claude/skills/informes-cartera-word")
    from estilo_cun import DocumentoCUN, AZUL_INST, TURQUESA, TIT, CUERPO

    D = DocumentoCUN()
    doc = D.doc
    par, h1, vineta, tabla = D.par, D.h1, D.vineta, D.tabla
    callout, kpi_row, nota, linea_acento = D.callout, D.kpi_row, D.nota, D.linea_acento

    D.portada("Cierre de Cartera - Agosto 2026",
              [("Dirigido a: ", "Oscar Penagos - Coordinacion de Recaudo y Cartera"),
               ("Periodo: ", "agosto de 2026")])
    ...
    D.guardar("Informe.docx")
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- Paleta institucional (seccion 1 de los lineamientos)
AZUL_MARINO = RGBColor(0x0C, 0x23, 0x40)   # titulos H1, encabezados de tabla
AZUL_INST = RGBColor(0x1B, 0x36, 0x5D)     # subtitulos H2
TURQUESA = RGBColor(0x00, 0x85, 0x9B)      # cifras KPI
VERDE_OSC = RGBColor(0x00, 0x7A, 0x33)     # vinetas
GRIS_INST = RGBColor(0x89, 0x8D, 0x8D)     # notas al pie, texto secundario
TEXTO = RGBColor(0x22, 0x22, 0x22)         # cuerpo
BLANCO = RGBColor(0xFF, 0xFF, 0xFF)

HEX_MARINO = "0C2340"
HEX_MANZANA = "84BD00"    # linea de acento, borde de callout
HEX_ZEBRA = "F8F9FA"      # filas pares de tabla
HEX_CALLOUT = "F0F4F8"    # fondo de callout
HEX_TURQUESA = "00859B"

# ---- Tipografia (seccion 2). Word sustituye si no estan instaladas.
TIT = "Montserrat"        # titulos, encabezados de tabla, cifras KPI
CUERPO = "Open Sans"      # cuerpo, tablas, notas

PIE_INSTITUCIONAL = ("Vicerrectoría de Servicios Digitales · Analítica financiera"
                     "   |   VIGILADA MINEDUCACIÓN")
ENCABEZADO_INSTITUCIONAL = "Corporación Unificada Nacional de Educación Superior - CUN"


class DocumentoCUN:
    """Documento Word con la identidad CUN ya aplicada.

    Los margenes por defecto (1,6 / 1,3 / 2,2 / 2,0 cm) son mas ajustados que los
    'ejecutivos' de los lineamientos a proposito: estos informes tienen tope duro
    de paginas y con los margenes anchos no caben. Si el informe no tiene tope,
    pasar margenes=(2.5, 2.5, 3.0, 2.5).
    """

    def __init__(self, margenes=(1.6, 1.3, 2.2, 2.0), con_encabezado=True):
        self.doc = Document()
        normal = self.doc.styles["Normal"]
        normal.font.name = CUERPO
        normal.font.size = Pt(10)
        normal.font.color.rgb = TEXTO
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), CUERPO)

        sup, inf, izq, der = margenes
        for s in self.doc.sections:
            s.top_margin = Cm(sup)
            s.bottom_margin = Cm(inf)
            s.left_margin = Cm(izq)
            s.right_margin = Cm(der)

        if con_encabezado:
            self._encabezado_y_pie()

    # ---------------------------------------------------------------- base
    @staticmethod
    def fuente(run, nombre, size, bold=False, italic=False, color=None):
        run.font.name = nombre
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        if color is not None:
            run.font.color.rgb = color
        run._element.rPr.rFonts.set(qn("w:eastAsia"), nombre)
        return run

    @staticmethod
    def shade(cell, hexc):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), hexc)
        tcPr.append(shd)

    @staticmethod
    def borde_izq(par, hexc, ancho=24):
        """Barra vertical de acento a la izquierda del parrafo (callout)."""
        pPr = par._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), str(ancho))
        left.set(qn("w:space"), "6")
        left.set(qn("w:color"), hexc)
        pbdr.append(left)
        pPr.append(pbdr)

    @staticmethod
    def fondo_par(par, hexc):
        pPr = par._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), hexc)
        pPr.append(shd)

    def _encabezado_y_pie(self):
        enc = self.doc.sections[0].header
        enc.is_linked_to_previous = False
        pe = enc.paragraphs[0]
        pe.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        self.fuente(pe.add_run(ENCABEZADO_INSTITUCIONAL), CUERPO, 8, color=GRIS_INST)

        pie = self.doc.sections[0].footer
        pie.is_linked_to_previous = False
        pp = pie.paragraphs[0]
        pp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self.fuente(pp.add_run(PIE_INSTITUCIONAL), CUERPO, 8, color=GRIS_INST)

    # ---------------------------------------------------------------- bloques
    def linea_acento(self, hexc=HEX_MANZANA, ancho=24, after=10):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(after)
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single")
        bot.set(qn("w:sz"), str(ancho))
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), hexc)
        pbdr.append(bot)
        pPr.append(pbdr)
        return p

    def par(self, texto="", size=10, bold=False, italic=False, color=TEXTO,
            fuente=CUERPO, after=5, before=0, align=None, interlineado=1.1):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.line_spacing = interlineado
        if align is not None:
            p.alignment = align
        if texto:
            self.fuente(p.add_run(texto), fuente, size, bold, italic, color)
        return p

    def h1(self, texto):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(7)
        p.paragraph_format.space_after = Pt(3)
        self.fuente(p.add_run(texto), TIT, 12.5, bold=True, color=AZUL_MARINO)
        return p

    def h2(self, texto, before=7):
        return self.par(texto, size=10, bold=True, color=AZUL_INST,
                        fuente=TIT, after=3, before=before)

    def vineta(self, texto, negrita_hasta=None):
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.55)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.08
        self.fuente(p.add_run("▪  "), CUERPO, 10, bold=True, color=VERDE_OSC)
        if negrita_hasta:
            self.fuente(p.add_run(negrita_hasta), CUERPO, 9.5, bold=True, color=AZUL_INST)
        self.fuente(p.add_run(texto), CUERPO, 9.5, color=TEXTO)
        return p

    def cell_text(self, cell, texto, bold=False, color=TEXTO, size=9,
                  align=None, fuente=CUERPO):
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(1)
        if align is not None:
            p.alignment = align
        self.fuente(p.add_run(str(texto)), fuente, size, bold, False, color)

    def tabla(self, headers, filas, anchos=None, size=8.5, alinear_der=None):
        """Tabla institucional: encabezado azul marino, zebra suave.

        OJO: `anchos` es una sugerencia. Word reajusta las columnas por su cuenta,
        asi que un encabezado largo se parte en dos lineas aunque le des ancho de
        sobra. La solucion que funciona es acortar el texto del encabezado
        ('% del total' -> '% total'), no ensanchar la columna.
        """
        alinear_der = alinear_der or []
        t = self.doc.add_table(rows=1, cols=len(headers))
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for j, hh in enumerate(headers):
            c = t.rows[0].cells[j]
            self.cell_text(c, hh, bold=True, color=BLANCO, size=size, fuente=TIT,
                           align=WD_ALIGN_PARAGRAPH.RIGHT if j in alinear_der else None)
            self.shade(c, HEX_MARINO)
        for i, fila in enumerate(filas):
            cells = t.add_row().cells
            for j, val in enumerate(fila):
                neg = j == 0 and str(val).startswith("Total")
                self.cell_text(cells[j], val, bold=neg, size=size,
                               align=WD_ALIGN_PARAGRAPH.RIGHT if j in alinear_der else None)
                if i % 2 == 1:
                    self.shade(cells[j], HEX_ZEBRA)
        if anchos:
            for row in t.rows:
                for j, w in enumerate(anchos):
                    row.cells[j].width = Cm(w)
        return t

    def callout(self, titulo, lineas, hexc_borde=HEX_MANZANA):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.15)
        self.borde_izq(p, hexc_borde)
        self.fondo_par(p, HEX_CALLOUT)
        self.fuente(p.add_run(titulo), TIT, 10.5, bold=True, color=AZUL_MARINO)
        for k, ln in enumerate(lineas):
            q = self.doc.add_paragraph()
            q.paragraph_format.space_before = Pt(0)
            q.paragraph_format.space_after = Pt(6 if k == len(lineas) - 1 else 1)
            q.paragraph_format.left_indent = Cm(0.15)
            q.paragraph_format.line_spacing = 1.08
            self.borde_izq(q, hexc_borde)
            self.fondo_par(q, HEX_CALLOUT)
            self.fuente(q.add_run(ln), CUERPO, 9.5, color=TEXTO)

    def kpi_row(self, items):
        """Fila de metricas destacadas: cifra grande turquesa + etiqueta gris.

        4 columnas es el maximo legible en carta. Con 5 las cifras se aprietan.
        """
        t = self.doc.add_table(rows=2, cols=len(items))
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for j, (cifra, etiqueta) in enumerate(items):
            self.cell_text(t.rows[0].cells[j], cifra, bold=True, color=TURQUESA,
                           size=15, align=WD_ALIGN_PARAGRAPH.CENTER, fuente=TIT)
            self.cell_text(t.rows[1].cells[j], etiqueta, color=GRIS_INST, size=8,
                           align=WD_ALIGN_PARAGRAPH.CENTER)
        return t

    def nota(self, texto):
        """Nota al pie de una seccion: fuente, metodologia o limitacion."""
        return self.par(texto, size=8.5, italic=True, color=GRIS_INST, after=3)

    def portada(self, titulo, metadatos, antetitulo="INFORME EJECUTIVO"):
        """Bloque de apertura: antetitulo, titulo, linea de acento y metadatos.

        `metadatos` es una lista de (etiqueta, valor); van en un solo parrafo
        separados por espacios, no en lineas sueltas, para no gastar alto.
        """
        self.par(antetitulo, size=9.5, bold=True, color=TURQUESA, fuente=TIT, after=1)
        self.par(titulo, size=19, bold=True, color=AZUL_MARINO, fuente=TIT, after=2)
        self.linea_acento(after=6)

        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        for etiqueta, valor in metadatos:
            self.fuente(p.add_run(etiqueta), CUERPO, 9, bold=True, color=AZUL_INST)
            self.fuente(p.add_run(valor + "     "), CUERPO, 9, color=TEXTO)
        return p

    def salto_pagina(self):
        self.doc.add_page_break()

    def guardar(self, ruta):
        self.doc.save(ruta)
        return ruta
