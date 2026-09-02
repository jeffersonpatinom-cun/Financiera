"""Informe ejecutivo (3 hojas) del cierre de agosto 2026.

Destinatario: Coordinacion de Recaudo y Cartera.
Bloques: cumplimiento de la meta, gestion del equipo de asesores, recaudo registrado.

Fuente de los datos: analisis_cierre_agosto_2026.sql sobre
  [Financiera].[Cartera_Meta_Comercial_Historico] + Snapshot_Mensual (bloque 1)
  [Financiera].[Cartera_CUN_Asesor_Unico]                            (bloques 2 y 3)

Reglas de calculo acordadas con la Coordinacion (2026-09-02):
  * Asesor_Unico excluyendo 'Reasignar en CRM' y 'Sin asignar'.
  * Gestion = Hora_modificacion_tipif en agosto (NO Hora_de_modificación).
  * Pago    = Fecha_de_pago en agosto, sumando Valor_pagado.
  * Q y marca de la meta SIEMPRE del snapshot 202609 (cierre de agosto).
  * El informe no nombra asesores: solo agregados del equipo.

Estilo: Lineamientos_Visuales_y_Comunicacion_CUN_Word.md
"""
import sys

sys.path.insert(0, ".claude/skills/informes-cartera-word")
from docx.shared import Pt                                       # noqa: E402
from estilo_cun import (DocumentoCUN, AZUL_MARINO, AZUL_INST,    # noqa: E402
                        TURQUESA, TEXTO, TIT, CUERPO)

sys.stdout.reconfigure(encoding="utf-8", errors="replace")

SALIDA = "Informe_Ejecutivo_Cierre_Agosto_2026.docx"

# La identidad visual CUN vive en el skill informes-cartera-word; aqui solo va el
# contenido. Los nombres cortos se rebindean para no reescribir cada llamada.
D = DocumentoCUN()
doc = D.doc
_fuente = D.fuente
par, h1, vineta, tabla = D.par, D.h1, D.vineta, D.tabla
callout, kpi_row, nota, linea_acento = D.callout, D.kpi_row, D.nota, D.linea_acento

# ================================================================== HOJA 1
par("INFORME EJECUTIVO", size=9.5, bold=True, color=TURQUESA, fuente=TIT, after=1)
par("Cierre de Cartera — Agosto 2026", size=19, bold=True,
    color=AZUL_MARINO, fuente=TIT, after=2)
linea_acento(after=6)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
_fuente(p.add_run("Dirigido a: "), CUERPO, 9, bold=True, color=AZUL_INST)
_fuente(p.add_run("Óscar Penagos — Coordinación de Recaudo y Cartera     "),
        CUERPO, 9, color=TEXTO)
_fuente(p.add_run("Periodo: "), CUERPO, 9, bold=True, color=AZUL_INST)
_fuente(p.add_run("agosto de 2026     "), CUERPO, 9, color=TEXTO)
_fuente(p.add_run("Elaboró: "), CUERPO, 9, bold=True, color=AZUL_INST)
_fuente(p.add_run("Analítica financiera CUN"), CUERPO, 9, color=TEXTO)

callout(
    "Conclusión",
    ["En agosto salieron de la meta 9.289 obligaciones de 6.248 estudiantes, equivalentes a "
     "$2.769,5 millones: un cumplimiento del 18,5% sobre las 50.146 obligaciones con que "
     "arrancó el mes. El equipo registró 61.767 gestiones sobre 18.716 personas y $5.483,9 "
     "millones en pagos.",
     "El hallazgo que exige decisión es de asignación, no de esfuerzo: 84.179 registros de "
     "28.619 personas permanecen sin asesor responsable y por tanto fuera de toda gestión."],
)

h1("1.  Cumplimiento de la meta")
kpi_row([("9.289", "OBLIGACIONES CUMPLIDAS"),
         ("6.248", "ESTUDIANTES"),
         ("$2.769", "MILLONES LIBERADOS"),
         ("18,5%", "CUMPLIMIENTO")])

par("La meta de agosto se cerró el 1 de septiembre. Comparamos el universo con el que se "
    "trabajó contra el que sobrevivió al corte: una obligación se considera cumplida cuando "
    "desaparece por completo de la cartera vigente.", before=6, after=4)

tabla(
    ["Concepto", "Obligaciones", "Estudiantes", "Saldo ($ MM)", "% total"],
    [["Universo con que arrancó agosto", "50.146", "20.748", "15.198,0", "100,0%"],
     ["Salieron por pago o normalización", "9.289", "6.248", "2.769,5", "18,2%"],
     ["Continúan en la meta de septiembre", "40.857", "15.005", "12.428,5", "81,8%"]],
    anchos=[7.0, 2.5, 2.4, 2.5, 2.0],
    alinear_der=[1, 2, 3, 4],
)

h1("2.  Cumplimiento por cuartil de gestión")
par("Los cuartiles se leen de la fotografía del cierre de agosto, no de la cartera de hoy: "
    "el proceso mensual recalcula la asignación Q en cada corrida y compararla contra la "
    "actual mezclaría dos criterios distintos.", after=4)
tabla(
    ["Cuartil", "Obligaciones", "Saldo ($ MM)", "Salieron", "% cumplimiento"],
    [["Q1 — mayor antigüedad", "7.007", "2.436,7", "515", "7,4%"],
     ["Q2", "7.565", "2.460,9", "725", "9,6%"],
     ["Q3", "6.955", "2.257,8", "777", "11,2%"],
     ["Q4 — vencimiento reciente", "28.619", "8.042,5", "7.272", "25,4%"]],
    anchos=[5.6, 2.7, 2.7, 2.4, 3.0],
    alinear_der=[1, 2, 3, 4],
)
par("", after=2)
vineta("el cumplimiento cae de forma sostenida a medida que aumenta la antigüedad: 25,4% en "
       "Q4 contra 7,4% en Q1. La obligación reciente se paga; la vencida hace más de un "
       "trimestre requiere una estrategia distinta a la llamada de seguimiento.",
       "Lectura principal: ")
nota("El cumplimiento se mide por salida completa de la obligación. No es posible medir el "
     "abono parcial de agosto: el refresco de saldos entró en operación el 1 de septiembre, "
     "de modo que los respaldos previos conservan el saldo de ingreso.")

# ================================================================== HOJA 2
doc.add_page_break()

h1("3.  Gestión del equipo de asesores")
kpi_row([("61.767", "GESTIONES"),
         ("18.716", "PERSONAS GESTIONADAS"),
         ("18 de 21", "ASESORES ACTIVOS"),
         ("32,1%", "EFECTIVIDAD")])

par("Contabilizamos como gestión cada cambio de tipificación registrado dentro del mes sobre "
    "cartera efectivamente asignada. La marca de modificación del sistema no sirve para esto: "
    "en agosto tocó el 93% de la base por una actualización masiva, que no es trabajo de "
    "asesor.", before=6, after=4)

tabla(
    ["Etapa del embudo", "Personas", "% sobre la etapa previa"],
    [["Personas en base asignada", "61.953", "—"],
     ["Gestionadas en agosto", "18.716", "30,2%"],
     ["Gestionadas que registraron pago", "6.011", "32,1%"]],
    anchos=[8.0, 3.0, 5.4],
    alinear_der=[1, 2],
)
par("", after=3)

par("Tipificaciones aplicadas en agosto", size=10, bold=True, color=AZUL_INST,
    fuente=TIT, after=3, before=5)
tabla(
    ["Tipificación", "Gestiones", "Personas", "% de la gestión"],
    [["Seguimiento al compromiso de pago", "34.020", "10.243", "55,1%"],
     ["No contesta", "15.677", "5.566", "25,4%"],
     ["Genera acuerdo de pago verbal", "6.647", "2.375", "10,8%"],
     ["Envío de correo tras llamada", "1.332", "652", "2,2%"],
     ["Confirma pago ya realizado", "2.504", "2.103", "4,1%"],
     ["No hay acuerdo de pago", "989", "452", "1,6%"]],
    anchos=[7.6, 2.6, 2.4, 3.8],
    alinear_der=[1, 2, 3],
)
par("", after=3)
vineta("la carga va de 33 a 10.958 gestiones por asesor, con un promedio de 3.431. Un solo "
       "asesor concentra el 17,7% de toda la actividad del mes, y tres de los 21 no "
       "registraron gestión alguna.", "Dispersión del esfuerzo: ")
vineta("una de cada cuatro gestiones (25,4%) termina en «no contesta». Es el mayor punto de "
       "fuga del proceso y apunta a calidad del dato de contacto, no a falta de trabajo.",
       "Contactabilidad efectiva: ")

h1("4.  Cumplimiento por perfil académico")
tabla(
    ["Marca académica", "Obligaciones", "Saldo ($ MM)", "Salieron", "% cumpl."],
    [["Sin registro de clase", "14.959", "4.355,0", "476", "3,2%"],
     ["Gestionable", "14.349", "5.217,1", "3.613", "25,2%"],
     ["Periodo en curso", "10.496", "2.489,2", "2.757", "26,3%"],
     ["Periodo perdido, prioridad alta", "5.955", "1.908,4", "562", "9,4%"],
     ["Periodo no ha iniciado", "3.980", "1.106,3", "1.804", "45,3%"]],
    anchos=[6.0, 2.5, 2.5, 2.2, 3.2],
    alinear_der=[1, 2, 3, 4],
)
vineta("«sin registro de clase» concentra $4.355,0 millones y solo cumplió el 3,2%. Es el "
       "grupo más grande de la meta y el que menos responde: sin vínculo académico activo, "
       "la gestión de cobro ordinaria no lo mueve.", "Segmento crítico: ")

# ================================================================== HOJA 3
doc.add_page_break()

h1("5.  Recaudo registrado por el equipo")
kpi_row([("19.002", "PAGOS REGISTRADOS"),
         ("11.804", "ESTUDIANTES"),
         ("$5.483", "MILLONES"),
         ("$288.642", "TICKET PROMEDIO")])

par("Las cifras de esta sección corresponden a lo que los asesores registraron en el CRM "
    "durante agosto sobre cartera asignada. No equivalen al recaudo institucional de caja, "
    "que se contabiliza por otra vía.", before=6, after=4)

par("Recaudo por periodo académico", size=10, bold=True, color=AZUL_INST,
    fuente=TIT, after=3)
tabla(
    ["Periodo", "Pagos", "Estudiantes", "Valor ($ MM)", "% del valor"],
    [["26V04", "3.316", "2.394", "1.007,3", "18,4%"],
     ["2026C", "1.454", "928", "820,8", "15,0%"],
     ["26V02", "2.750", "1.595", "536,5", "9,8%"],
     ["26V03", "2.666", "1.899", "493,0", "9,0%"],
     ["26ES4", "1.326", "976", "480,4", "8,8%"],
     ["26ES3", "1.461", "1.038", "417,1", "7,6%"],
     ["Resto de periodos", "6.029", "—", "1.728,8", "31,4%"]],
    anchos=[3.6, 2.6, 2.9, 2.9, 2.4],
    alinear_der=[1, 2, 3, 4],
)
par("", after=3)

par("Distribución del ticket", size=10, bold=True, color=AZUL_INST, fuente=TIT, after=3)
tabla(
    ["Rango", "Bajo 100k", "100k – 300k", "300k – 600k", "600k – 1M", "Más de 1M"],
    [["Pagos", "3.888", "8.261", "5.695", "722", "433"],
     ["Valor ($ MM)", "61,7", "1.669,2", "2.269,9", "521,5", "961,6"]],
    anchos=[2.8, 2.7, 2.6, 2.6, 2.4, 2.4],
    alinear_der=[1, 2, 3, 4, 5],
)
par("", after=3)
vineta("el rango de 300 a 600 mil pesos aporta $2.269,9 millones, el 41,4% del valor con el "
       "30% de los pagos. Es la cuota típica del crédito CLTIENE y donde rinde más el "
       "esfuerzo de contacto.", "Concentración del valor: ")

h1("6.  Observaciones de la Analítica")
vineta("84.179 registros de 28.619 personas están marcados «reasignar en CRM» o «sin "
       "asignar». Es el 27% de la cartera del CRM y ningún asesor la tiene a cargo, de modo "
       "que queda fuera de cualquier indicador de gestión. Recomendamos resolver la "
       "asignación antes del cierre de septiembre.", "Cartera sin responsable: ")
vineta("los campos medio de pago, tipo de cartera y regional están sin diligenciar en más "
       "del 89% de los registros de pago, lo que impide analizar el recaudo por canal. Es "
       "un ajuste de captura en el CRM, no de la analítica.", "Campos sin diligenciar: ")
vineta("la tabla del CRM es una fotografía del día y se reconstruye en cada corrida. Lo que "
       "se tipificó en agosto y se sobrescribió después ya no es visible, por lo que las "
       "cifras de gestión son un piso y no un techo.", "Ausencia de histórico: ")

callout(
    "Recomendación",
    ["Reasignar los 84.179 registros sin responsable y atacar el segmento «sin registro de "
     "clase» con una estrategia distinta a la llamada de seguimiento: concentra $4.355,0 "
     "millones y solo respondió el 3,2%.",
     "Sobre el 25,4% de gestiones que terminan en «no contesta», priorizar la depuración del "
     "dato de contacto antes de aumentar el volumen de llamadas."],
    hexc_borde="00859B",
)

par("", after=2)
nota("Fuentes: [Financiera].[Cartera_Meta_Comercial_Historico] y su snapshot de cierre, y "
     "[Financiera].[Cartera_CUN_Asesor_Unico]. Extracción del 2 de septiembre de 2026. "
     "Cifras en millones de pesos colombianos.")

D.guardar(SALIDA)
print(f"OK -> {SALIDA}")
