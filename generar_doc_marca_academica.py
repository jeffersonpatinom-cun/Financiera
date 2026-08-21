"""Genera el documento Word con la logica de negocio de MARCA_ACADEMICA.

Cifras: ejecucion real de marca_academica_combinaciones.sql contra
CUN_REPOSITORIO.FINANCIERA.Cartera_Gestion.
  - comb.json          -> 3 resultados del query (matriz, resumen, nivel cliente unico)
  - comb_detalle.json  -> clientes distintos por marca de detalle
Todos los conteos del documento son a nivel de CLIENTES UNICOS.
"""
import json
import os
import sys

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout.reconfigure(encoding="utf-8", errors="replace")

AZUL = RGBColor(0x1F, 0x4E, 0x78)
AZUL2 = RGBColor(0x2E, 0x75, 0xB6)
GRIS = RGBColor(0x59, 0x59, 0x59)
ROJO = RGBColor(0xC0, 0x00, 0x00)
VERDE = RGBColor(0x38, 0x76, 0x1D)
NARANJA = RGBColor(0xC5, 0x5A, 0x11)
MORADO = RGBColor(0x7C, 0x3A, 0x8C)

BASE = os.path.dirname(os.path.abspath(__file__))
SALIDA = os.path.join(BASE, "Logica_MARCA_ACADEMICA_Recuperacion_Cartera.docx")

TOT_OBLIG = 257849
TOT_CLIENTES = 81753

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
for s in doc.sections:
    s.top_margin = Cm(1.8); s.bottom_margin = Cm(1.8)
    s.left_margin = Cm(2.0); s.right_margin = Cm(2.0)


# ------------------------------------------------------------------ helpers
def shade(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), hexc)
    tcPr.append(shd)


def cell_text(cell, texto, bold=False, color=None, size=8.5, italic=False, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    for k, part in enumerate(str(texto).split("\n")):
        if k:
            p = cell.add_paragraph()
        if align is not None:
            p.alignment = align
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(part); r.bold = bold; r.italic = italic; r.font.size = Pt(size)
        if color is not None:
            r.font.color.rgb = color


def par(texto="", bold=False, color=None, size=10.5, italic=False, after=6, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(after)
    if texto:
        r = p.add_run(texto); r.bold = bold; r.italic = italic; r.font.size = Pt(size)
        if color is not None:
            r.font.color.rgb = color
    return p


def mono(texto, size=8.5, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.6)
    r = p.add_run(texto)
    r.font.name = "Consolas"; r.font.size = Pt(size)
    r.font.color.rgb = color if color is not None else GRIS
    return p


def h1(texto):
    doc.add_paragraph()
    h = doc.add_heading(level=1); r = h.add_run(texto)
    r.font.color.rgb = AZUL; r.font.size = Pt(15)


def h2(texto):
    h = doc.add_heading(level=2); r = h.add_run(texto)
    r.font.color.rgb = AZUL2; r.font.size = Pt(12)


def tabla(encabezados, filas, anchos=None, size=8.5, header_fill="1F4E78", num_cols=()):
    t = doc.add_table(rows=1, cols=len(encabezados))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, e in enumerate(encabezados):
        c = t.rows[0].cells[i]
        shade(c, header_fill)
        cell_text(c, e, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=size)
    for j, fila in enumerate(filas):
        cells = t.add_row().cells
        for i, v in enumerate(fila):
            cell_text(cells[i], v, size=size,
                      align=WD_ALIGN_PARAGRAPH.RIGHT if i in num_cols else None)
            if j % 2 == 1:
                shade(cells[i], "F2F6FA")
    if anchos:
        for row in t.rows:
            for i, w in enumerate(anchos):
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def fmt(n):
    return f"{int(n):,}".replace(",", ".")


def pct(n, tot=None):
    tot = tot or TOT_CLIENTES
    return f"{100.0*int(n)/tot:.2f} %".replace(".", ",")


# ------------------------------------------------------------------ datos
with open(os.path.join(BASE, "comb.json"), encoding="utf-8") as f:
    SETS = json.load(f)
MATRIZ = SETS[0]["rows"]
RESUMEN = SETS[1]["rows"]
UNICOS = SETS[2]["rows"]

with open(os.path.join(BASE, "comb_detalle.json"), encoding="utf-8") as f:
    DETALLE = json.load(f)          # [marca, detalle, clientes, obligaciones]

cli_marca = {r["MARCA_ACADEMICA"]: int(r["CLIENTES_UNICOS"]) for r in RESUMEN}
obl_marca = {r["MARCA_ACADEMICA"]: int(r["OBLIGACIONES"]) for r in RESUMEN}
uni_marca = {r["MARCA_ACADEMICA"]: int(r["CLIENTES_UNICOS"]) for r in UNICOS}
det_cli = {d[1]: int(d[2]) for d in DETALLE}
det_obl = {d[1]: int(d[3]) for d in DETALLE}
det_marca = {d[1]: d[0] for d in DETALLE}


def combos_de(marca, n=6):
    return sorted([c for c in MATRIZ if c["MARCA_ACADEMICA"] == marca],
                  key=lambda x: -int(x["CLIENTES_UNICOS"]))[:n]


# ================================================================== PORTADA
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("MARCA_ACADEMICA")
r.font.size = Pt(26); r.bold = True; r.font.color.rgb = AZUL

t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Logica de segmentacion academico-crediticia para la gestion de recuperacion de cartera")
r.font.size = Pt(12.5); r.font.color.rgb = AZUL2

t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Analitica financiera - Universidad CUN")
r.font.size = Pt(10.5); r.font.color.rgb = GRIS

t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run(f"Fuente: CUN_REPOSITORIO.FINANCIERA.Cartera_Gestion   |   {fmt(TOT_CLIENTES)} clientes unicos   |   "
              f"{fmt(TOT_OBLIG)} obligaciones   |   {len(MATRIZ)} casuisticas observadas")
r.font.size = Pt(9); r.italic = True; r.font.color.rgb = GRIS

doc.add_paragraph()

# ================================================================== 1
h1("1. Para que sirve la marca")
par("MARCA_ACADEMICA responde una sola pregunta operativa: ¿que le pasa hoy a este deudor con la "
    "institucion, y por lo tanto que tipo de gestion de cobro corresponde? No es un dato academico: "
    "es un enrutador de la estrategia de recuperacion.")
par("El principio de fondo en recuperacion de cartera educativa es que la deuda de un estudiante que "
    "sigue estudiando NO se cobra igual que la de uno que ya perdio el vinculo. En el primer caso la "
    "matricula del proximo periodo es la palanca (retencion y acuerdo de pago); en el segundo esa "
    "palanca no existe y el caso migra a cobro persuasivo, prejuridico o castigo. Marcar mal a un "
    "estudiante activo como moroso duro destruye matricula; marcar como gestionable a un desertor con "
    "riesgo crediticio adverso pierde tiempo de asesor y valor recuperable.")
par("Variables que construyen la marca, todas disponibles hoy en Cartera_Gestion:", bold=True)
tabla(
    ["Columna fuente", "Que aporta a la decision"],
    [
        ["NOMBRE_TIPO_CLIENTE", "Segmento del deudor (ESTUDIANTES / COMERCIAL / COLABORADORES). Separa a quien no tiene vida academica que evaluar."],
        ["ESTADO", "Estado de calendario del periodo facturado (ACTIVO / NO ACTIVO / PERIODO NO HA INICIADO). Define si el hecho academico ya ocurrio o todavia esta abierto."],
        ["ESTADO_ALUMNO", "Vinculo del deudor con la institucion (1-Activo / 2-Egresado / 3-Graduado). Define si queda palanca de matricula."],
        ["RES_PERFIL_RIESGO", "Perfil crediticio externo del deudor. Define la probabilidad de pago voluntario y la urgencia de escalar."],
        ["ultimoaccesoplataformlimpio", "Ultimo acceso a plataforma. Es la unica evidencia de vida academica cuando la nota todavia no esta cargada."],
        ["PROMEDIO", "Resultado academico del periodo. Es la evidencia dura de aprobacion o perdida."],
    ],
    anchos=[5.0, 12.0])

h2("1.1 Como leer los numeros de este documento")
par(f"Cartera_Gestion tiene una fila por OBLIGACION, no por persona: {fmt(TOT_OBLIG)} obligaciones "
    f"que corresponden a {fmt(TOT_CLIENTES)} clientes unicos, es decir un promedio de 3,2 "
    "obligaciones por deudor (distintos periodos y creditos de la misma persona).")
par("Todos los conteos de este documento estan a nivel de CLIENTES UNICOS, que es la unidad con la "
    "que se dimensiona una campaña: lo que se contacta es personas, no filas. Contar filas infla los "
    "volumenes hasta tres veces y lleva a sobredimensionar la capacidad de gestion requerida.",
    bold=True)
par("Una misma persona puede tener obligaciones en marcas distintas (por ejemplo un periodo perdido "
    "y otro en curso). Para la cola de trabajo se le asigna UNA sola marca: la de mayor urgencia. Esa "
    "es la vista que se presenta en el resumen final y la que suma exactamente el universo de la "
    "cartera.", italic=True, size=9.5, color=GRIS)

# ================================================================== 2
h1("2. Tipificacion de cada variable")
par("Antes de decidir la marca, cada variable se reduce a un dominio cerrado y sin vacios. Un dato "
    "ausente no es una categoria: obliga a decidir que significa esa ausencia, y esa decision es "
    "parte de la regla de negocio.")

h2("2.1 NOMBRE_TIPO_CLIENTE -> segmento")
tabla(["Valor", "Clientes", "Obligaciones", "Lectura de negocio"],
      [["ESTUDIANTES", fmt(81712), fmt(248418), "Unico segmento con vida academica evaluable."],
       ["COMERCIAL", fmt(40), fmt(9426), "Convenios y empresas. Sin notas, sin plataforma, sin periodo academico."],
       ["COLABORADORES", fmt(1), fmt(5), "Volumen marginal; se agrupa con el segmento no estudiantil."]],
      anchos=[3.6, 2.4, 2.8, 8.2], num_cols=(1, 2))

h2("2.2 ESTADO -> ESTADO_PERIODO")
tabla(["Valor", "Obligaciones", "Tipificacion", "Lectura de negocio"],
      [["ACTIVO", fmt(105560), "ACTIVO", "Periodo abierto: el resultado academico aun no esta escrito."],
       ["NO ACTIVO", fmt(145074), "NO ACTIVO", "Periodo cerrado: lo que diga la nota es definitivo."],
       ["PERIODO NO HA INICIADO", fmt(6379), "PERIODO NO HA INICIADO", "Facturado por anticipado; no hay hecho academico que evaluar."],
       ["(vacio)", fmt(836), "SIN DATO", "Se preserva como categoria propia: no se asume ni abierto ni cerrado."]],
      anchos=[4.2, 2.4, 3.6, 6.8], num_cols=(1,))

h2("2.3 PROMEDIO -> [ESTADO NOTAS]")
par("Corte en 3.0, la nota minima aprobatoria del reglamento academico. La ausencia de nota NO se "
    "interpreta como perdida: en un periodo abierto simplemente todavia no se ha cargado.")
tabla(["Regla", "Obligaciones", "Valor", "Lectura de negocio"],
      [["PROMEDIO sin dato", fmt(143874), "SIN NOTA", "Sin evidencia academica. Se resuelve con el estado del periodo y el acceso a plataforma."],
       ["PROMEDIO >= 3.0", fmt(43508), "APROBO", "Cumplio. Sigue siendo cliente: el cobro debe cuidar la relacion."],
       ["PROMEDIO < 3.0", fmt(70467), "PERDIO", "Fallo academico. Alta probabilidad de desercion y de deterioro de la deuda."]],
      anchos=[3.8, 2.4, 2.4, 8.4], num_cols=(1,))

h2("2.4 ultimoaccesoplataformlimpio -> [acceso plataforma]")
par("Es un desempate, no un criterio principal. Sirve exactamente para el hueco de 143.874 "
    "obligaciones sin nota: si el deudor entro a la plataforma, sigue conectado academicamente aunque "
    "el sistema aun no haya cargado calificaciones. Sin acceso y sin nota, no hay evidencia de vida "
    "academica.")
tabla(["Regla", "Obligaciones", "Valor"],
      [["Registra acceso a plataforma", fmt(219288), "SI"],
       ["No registra acceso", fmt(38561), "NO"]],
      anchos=[8.0, 3.0, 2.5], num_cols=(1,))

# ================================================================== 3
h1("3. RES_PERFIL_RIESGO: por que el SI/NO no es presencia del dato")
par("Este es el punto que mas cambia el resultado de la marca.", bold=True)
par("La escala de RES_PERFIL_RIESGO es ordinal y viene acompañada de un puntaje. Verificado contra la "
    "cartera, los cortes de la escala son exactos y no se solapan:")
tabla(["RES_PERFIL_RIESGO", "Puntaje observado", "Obligaciones", "% cartera", "Interpretacion de cobranza"],
      [["Riesgo Excelente", "904 - 1000", fmt(671), "0,26 %", "Paga solo. Un recordatorio basta."],
       ["Riesgo Muy Bueno", "800 - 899", fmt(5126), "1,99 %", "Buen pagador. Gestion amable."],
       ["Riesgo Bueno", "670 - 798", fmt(13321), "5,17 %", "Pagador estandar. Cobro ordinario."],
       ["Riesgo Regular", "580 - 668", fmt(9893), "3,84 %", "Fragil. Requiere acuerdo de pago dirigido."],
       ["Riesgo Alto", "0 - 579", fmt(56289), "21,83 %", "Deterioro probable. Escalar y priorizar."],
       ["Sin perfil", "-", fmt(172549), "66,92 %", "No evaluado. Ausencia de dato, no ausencia de riesgo."]],
      anchos=[3.6, 3.0, 2.4, 2.2, 5.8], num_cols=(2, 3))

par("La regla original marcaba SI cuando el perfil simplemente existia. Eso confunde dos cosas "
    "distintas: tener el dato y tener mal perfil. Bajo esa regla, los deudores con perfil Bueno, Muy "
    "Bueno o Excelente -- justamente los de mayor probabilidad de pago voluntario -- entraban a "
    "PRIORIDAD ALTA solo porque la consulta de buro devolvio resultado. El efecto practico es doble: "
    "se gasta el recurso escaso (asesor de cobro persuasivo) en quien iba a pagar solo, y se diluye "
    "el sentido de la cola de alta prioridad.")
par("Criterio adoptado:", bold=True)
mono("RIESGO = 'SI'  <=>  perfil Riesgo Alto o Riesgo Regular   (puntaje inferior a 670)\n"
     "RIESGO = 'NO'  <=>  Bueno / Muy Bueno / Excelente, o sin perfil", color=AZUL)
par("El corte en 670 no es arbitrario: es la frontera que la propia escala pone entre Riesgo Regular "
    "(hasta 668) y Riesgo Bueno (desde 670), y coincide con el umbral de uso general en riesgo de "
    "credito. SI significa riesgo adverso, no 'tiene perfil'.")
par("La ausencia de perfil se trata como NO riesgo. Con dos tercios de la cartera sin evaluar, "
    "asumir riesgo por defecto marcaria de alta prioridad a la mayoria del portafolio y la "
    "priorizacion dejaria de discriminar. La cobertura del dato se conserva en una columna aparte, "
    "COBERTURA_PERFIL, que es de diagnostico y no interviene en la marca.", italic=True)
par("Impacto medido del cambio (misma escalera, solo cambia el SI/NO):", bold=True)
tabla(["Marca con regla anterior", "Marca con riesgo adverso", "Obligaciones"],
      [["PERIODO PERDIDO, PRIORIDAD ALTA", "GESTIONABLE", fmt(4117)],
       ["PERIODO PERDIDO, PRIORIDAD ALTA", "SIN REGISTRO DE CLASE", fmt(405)],
       ["(sin cambio)", "(sin cambio)", fmt(253327)]],
      anchos=[6.0, 6.0, 3.0], num_cols=(2,))
par("4.522 obligaciones salen de la cola de alta prioridad. Son deudores con perfil crediticio "
    "favorable que estaban ocupando capacidad de gestion intensiva sin justificacion de riesgo.",
    size=9.5, italic=True, color=GRIS)

# ================================================================== 4
h1("4. La escalera de decision")
par("La marca se asigna con una escalera excluyente: la primera condicion que se cumple gana y las "
    "demas no se evaluan. El orden no es estetico, es la jerarquia de evidencia:", bold=True)
mono("0. Segmento          ->  ¿es un deudor con vida academica?\n"
     "1. Calendario        ->  ¿ya ocurrio el hecho academico?\n"
     "2. Periodo abierto   ->  si sigue en curso, la nota parcial no condena\n"
     "3. Fallo o riesgo    ->  evidencia dura de deterioro\n"
     "4. Cumplimiento      ->  aprobo o cerro ciclo\n"
     "5. Catch-all         ->  sin evidencia de ningun tipo", color=AZUL)
par("Consecuencia de diseño: la escalera nunca deja un registro sin marca. Toda obligacion cae en "
    "exactamente una de las seis marcas, lo que permite que el tablero sume el 100 % de la cartera "
    "sin una categoria residual invisible.", italic=True, size=9.5)

BLOQUES = [
    ("CARTERA EMPRESARIAL", 0, MORADO,
     "El deudor no es un estudiante: es una empresa, un convenio o un colaborador.",
     "NOMBRE_TIPO_CLIENTE <> 'ESTUDIANTES'",
     "Se evalua de PRIMERO, antes que cualquier variable academica. Un NIT empresarial no tiene notas, "
     "no entra a plataforma y no cursa periodos: aplicarle la lectura academica produce resultados sin "
     "sentido. Hoy, sin esta separacion, estos deudores se dispersan por las marcas academicas: la "
     "mayoria cae en 'sin registro de clase' -- que para una empresa no significa nada -- y cerca de "
     "550 obligaciones se filtran a la cola de retencion estudiantil, donde el mensaje que reciben no "
     "les corresponde. Son pocos clientes pero mucha obligacion: 41 deudores concentran 9.431 "
     "obligaciones, un promedio de 230 cada uno, de modo que distorsionan cualquier indicador que se "
     "mida por filas.",
     "Ruta de cobro corporativa: interlocucion con el area responsable del convenio, conciliacion de "
     "saldos y acuerdos institucionales. No aplica gestion academica ni mensajes de retencion."),
    ("PERIODO NO HA INICIADO", 1, VERDE,
     "El periodo facturado todavia no arranca por calendario.",
     "ESTADO_PERIODO = 'PERIODO NO HA INICIADO'",
     "Ninguna otra variable tiene sentido todavia: no hay nota que exista ni acceso que esperar. "
     "Cualquier lectura academica en este estado seria un artefacto. La obligacion existe (matricula "
     "facturada) pero el servicio aun no se presta.",
     "Cobro preventivo y recordatorio de fecha limite. NO es mora academica y no debe entrar a colas "
     "de cobro persuasivo. Es tambien la ventana de mayor rendimiento: cobrar antes de que inicien "
     "clases evita el problema entero."),
    ("PERIODO EN CURSO", 2, AZUL2,
     "El periodo esta abierto: el resultado academico todavia se puede cambiar.",
     "ESTADO_PERIODO = 'ACTIVO'\n  O BIEN (registra acceso a plataforma Y no perdio Y el periodo no esta cerrado)",
     "Dos caminos entran aqui. El primero es el calendario: si el periodo esta activo, la nota es "
     "parcial y no constituye fallo consumado. El segundo cubre el hueco de datos: cuando el estado "
     "del periodo no viene informado pero el deudor accedio a plataforma y no tiene nota perdida, se "
     "le da el beneficio de la duda academica. La condicion de que el periodo no este cerrado impide "
     "que ese beneficio se extienda a periodos ya terminados, donde un acceso a plataforma es solo un "
     "rastro historico.",
     "Gestion de retencion, no de cobranza dura. El deudor todavia es estudiante y la matricula del "
     "periodo siguiente es la palanca real. El mensaje es 'no pierdas tu cupo', no 'estas en mora'. "
     "Escalar aqui produce desercion y convierte una cartera recuperable en una incobrable."),
    ("PERIODO PERDIDO, PRIORIDAD ALTA", 3, ROJO,
     "Hay evidencia dura de deterioro: fallo academico consumado o perfil crediticio adverso.",
     "[ESTADO NOTAS] = 'PERDIO'\n  O BIEN (riesgo adverso Y vinculo vigente/egresado/graduado)",
     "Se llega por dos rutas independientes y ambas justifican la misma urgencia. La ruta academica: "
     "perdio el periodo, la probabilidad de que no se rematricule es alta y con ella se pierde la "
     "unica palanca blanda de cobro. La ruta crediticia: aunque haya aprobado, un perfil Alto o "
     "Regular anticipa incumplimiento y conviene actuar antes de que la mora escale de tramo. El "
     "filtro por vinculo evita escalar registros sin identificacion academica confiable.",
     "Cola de maxima intensidad: contacto telefonico prioritario, acuerdo de pago con cuota inicial y "
     "evaluacion temprana de traslado a prejuridico. Es la cola donde la capacidad de asesores debe "
     "concentrarse."),
    ("GESTIONABLE", 4, NARANJA,
     "Cumplio academicamente o ya cerro su ciclo, y no muestra riesgo crediticio adverso.",
     "[ESTADO NOTAS] = 'APROBO'\n  O BIEN vinculo Egresado / Graduado",
     "Deuda sana en cabeza de alguien con historial de cumplimiento. El egresado y el graduado entran "
     "aqui aunque no tengan nota del periodo: ya terminaron, no hay hecho academico pendiente, y "
     "conservan un incentivo fuerte y concreto -- la institucion retiene grados y certificaciones "
     "hasta el paz y salvo. Ese incentivo es mas efectivo que cualquier escalamiento.",
     "Cobro ordinario y automatizable: recordatorio por correo, WhatsApp o SMS, y en el caso de "
     "egresados el recordatorio del tramite de grado pendiente. Bajo costo por peso recuperado; no "
     "consume asesor."),
    ("SIN REGISTRO DE CLASE", 5, GRIS,
     "Periodo cerrado sin ninguna evidencia academica: ni nota, ni acceso, ni cierre de ciclo.",
     "Ninguna condicion anterior se cumple (catch-all)",
     "Es el residuo deliberado de la escalera y merece leerse como diagnostico, no como categoria de "
     "gestion. Este perfil corresponde tipicamente a matricula facturada con desercion temprana: se "
     "genero la obligacion, el estudiante nunca activo su vida academica y el periodo cerro. Tambien "
     "recoge fallas de cobertura del cruce academico. Que sea catch-all garantiza que ninguna "
     "obligacion quede sin marca, pero obliga a vigilar su tamaño: si crece, casi siempre es un "
     "problema de datos, no de negocio.",
     "Verificacion previa a la gestion: confirmar que la obligacion es exigible y que los datos de "
     "contacto sirven. Los casos con acceso a plataforma registrado son recuperables por contacto "
     "directo; los que no tienen ni acceso ni nota son candidatos naturales a depuracion, castigo o "
     "cobro juridico segun cuantia y antiguedad."),
]

for nombre, num, color, resumen, regla, porque, accion in BLOQUES:
    h2(f"4.{num}  {nombre}   ({fmt(cli_marca[nombre])} clientes - {fmt(obl_marca[nombre])} obligaciones)")
    par(resumen, bold=True, color=color)
    par("Regla:", bold=True, size=9.5, after=2)
    mono(regla)
    par("Por que se define asi:", bold=True, size=9.5, after=2)
    par(porque, size=10)
    par("Accion de cobranza:", bold=True, size=9.5, after=2)
    par(accion, size=10)
    par("Casuisticas principales que caen aqui:", bold=True, size=9.5, after=2)
    tabla(["ESTADO_PERIODO", "VINCULO", "RIESGO", "ACCESO", "NOTAS", "Clientes"],
          [[c["ESTADO_PERIODO"], c["ESTADO_ALUMNO"], c["RES_PERFIL_RIESGO"],
            c["acceso plataforma"], c["ESTADO NOTAS"], fmt(c["CLIENTES_UNICOS"])]
           for c in combos_de(nombre)],
          anchos=[4.0, 2.6, 1.8, 1.8, 2.2, 2.2], size=8, num_cols=(5,))

# ================================================================== 5
h1("5. MARCA_ACADEMICA_DETALLE: abrir la marca sin romper el tablero")
par("Las seis marcas anteriores son el contrato con el tablero y no se subdividen alli. Pero dentro "
    "de una misma marca conviven casos que no se gestionan igual: en PERIODO PERDIDO, PRIORIDAD ALTA "
    "hay quien perdio el periodo, quien tiene riesgo crediticio y quien tiene las dos cosas. La "
    "columna MARCA_ACADEMICA_DETALLE separa esos casos con las mismas variables y permite ordenar la "
    "cola de trabajo del asesor.")
ACCIONES_DET = {
    "PERIODO PERDIDO + RIESGO CREDITICIO": "Doble deterioro. Primera prioridad absoluta de la cola.",
    "PERIODO PERDIDO": "Fallo academico sin senal crediticia adversa. Acuerdo de pago y retencion.",
    "RIESGO CREDITICIO ADVERSO": "Aprobo pero el buro advierte. Contacto preventivo antes de que escale la mora.",
    "PERIODO EN CURSO - RIESGO CREDITICIO": "Estudiante activo con perfil fragil. Ofrecer financiacion antes de la mora.",
    "PERIODO EN CURSO": "Retencion y recordatorio suave.",
    "PERIODO NO HA INICIADO": "Cobro preventivo anticipado.",
    "GESTIONABLE": "Campaña masiva automatizada.",
    "SIN REGISTRO DE CLASE - CON CONEXION": "Hay rastro de conexion: contacto directo, es recuperable.",
    "SIN REGISTRO DE CLASE - SIN CONTACTO": "Sin ningun rastro: verificar exigibilidad, depurar o escalar a juridico.",
    "CARTERA EMPRESARIAL - COMERCIAL": "Interlocucion corporativa y conciliacion de saldos.",
    "CARTERA EMPRESARIAL - COLABORADORES": "Tramite interno con Gestion Humana.",
}
tabla(["MARCA_ACADEMICA_DETALLE", "Clientes", "Oblig.", "Que hacer"],
      [[d, fmt(det_cli[d]), fmt(det_obl[d]), ACCIONES_DET.get(d, "")]
       for d, _ in sorted(det_cli.items(), key=lambda x: -x[1])],
      anchos=[5.6, 1.8, 1.8, 8.0], size=8.5, num_cols=(1, 2))
par(f"El foco de cualquier campaña de recuperacion son los "
    f"{fmt(det_cli.get('PERIODO PERDIDO + RIESGO CREDITICIO', 0))} clientes que perdieron el periodo "
    "y ademas tienen perfil crediticio adverso: es el nucleo del deterioro de la cartera.",
    size=9.5, italic=True, color=GRIS)
par("Nota: la subcategoria 'PERIODO EN CURSO - ALERTA ACADEMICA' esta prevista en la logica pero hoy "
    "no tiene registros, porque los periodos activos aun no traen notas cargadas. Se deja definida "
    "para cuando el cargue academico se haga durante el periodo.", size=9, italic=True, color=GRIS)

# ================================================================== 6
h1("6. Casos borde y advertencias de lectura")
tabla(["Situacion", "Como la resuelve la logica", "Riesgo si se ignora"],
      [["Estado de periodo sin dato (836 obligaciones)",
        "Se tipifica como SIN DATO. Puede llegar a PERIODO EN CURSO solo si hay acceso a plataforma y no perdio; nunca por calendario.",
        "Asumirlo cerrado escalaria a estudiantes vigentes."],
       ["Vinculo academico sin dato (12.936 obligaciones)",
        "No habilita la ruta crediticia de la prioridad alta ni la ruta de egresado de GESTIONABLE.",
        "Una parte son deudores no estudiantiles; escalarlos como estudiantes distorsiona la cola."],
       ["Nota perdida en periodo activo",
        "El calendario gana: queda PERIODO EN CURSO, y el detalle lo separa como alerta academica.",
        "Tratar una nota parcial como fallo consumado dispara cobro duro sobre estudiantes que aun pueden recuperarse."],
       ["Dos tercios de la cartera sin perfil crediticio",
        "Se trata como NO riesgo. La cobertura se mide aparte, sin intervenir en la marca.",
        "Asumir riesgo por defecto marcaria de alta prioridad a la mayoria de la cartera y la priorizacion dejaria de servir."],
       ["41 clientes empresariales con 9.431 obligaciones",
        "Se separan en CARTERA EMPRESARIAL antes de cualquier lectura academica.",
        "Un promedio de 230 obligaciones por cliente distorsiona todo indicador medido por filas."],
       ["Un mismo deudor en varias marcas",
        "Para la cola de trabajo se le asigna la marca de mayor urgencia (ver resumen final).",
        "Contactarlo por dos vias con mensajes contradictorios."],
       ["Crecimiento de SIN REGISTRO DE CLASE",
        "Es catch-all por diseño; su tamaño es un indicador de calidad del cruce academico.",
        "Un salto subito casi nunca es negocio: suele ser una fuente academica que dejo de cruzar."]],
      anchos=[3.8, 7.0, 6.2], size=8.5)

# ================================================================== 7
h1("7. Resumen: como quedaria tipificada la cartera")
par("Esta es la fotografia que produce la logica propuesta sobre la cartera actual.", bold=True)

h2("7.1 Cola de gestion: clientes unicos, una sola marca por persona")
par(f"Cada uno de los {fmt(TOT_CLIENTES)} deudores aparece UNA sola vez, en su marca de mayor "
    "urgencia. Esta es la tabla con la que se dimensiona la operacion y se reparte la capacidad de "
    "los asesores.", size=10)
ORDEN = ["PERIODO PERDIDO, PRIORIDAD ALTA", "PERIODO EN CURSO", "GESTIONABLE",
         "SIN REGISTRO DE CLASE", "PERIODO NO HA INICIADO", "CARTERA EMPRESARIAL"]
INTENSIDAD = {
    "PERIODO PERDIDO, PRIORIDAD ALTA": "Alta - asesor dedicado",
    "PERIODO EN CURSO": "Media - retencion",
    "GESTIONABLE": "Baja - campaña masiva",
    "SIN REGISTRO DE CLASE": "Verificacion previa",
    "PERIODO NO HA INICIADO": "Preventiva",
    "CARTERA EMPRESARIAL": "Ruta corporativa",
}
tabla(["MARCA_ACADEMICA", "Clientes unicos", "% cartera", "Intensidad de gestion"],
      [[m, fmt(uni_marca[m]), pct(uni_marca[m]), INTENSIDAD[m]] for m in ORDEN]
      + [["TOTAL", fmt(sum(uni_marca.values())), "100,00 %", ""]],
      anchos=[6.0, 3.0, 2.4, 5.6], size=9, num_cols=(1, 2))

h2("7.2 Volumen total por marca (obligaciones y clientes)")
par("Aqui una misma persona puede contarse en mas de una marca, porque tiene obligaciones de "
    "periodos distintos. Sirve para dimensionar saldo y esfuerzo de conciliacion, no para "
    "dimensionar contactos.", size=10)
tabla(["MARCA_ACADEMICA", "Clientes", "Obligaciones", "Oblig. x cliente"],
      [[m, fmt(cli_marca[m]), fmt(obl_marca[m]),
        f"{obl_marca[m]/cli_marca[m]:.1f}".replace(".", ",")] for m in ORDEN]
      + [["TOTAL", fmt(TOT_CLIENTES), fmt(TOT_OBLIG), "3,2"]],
      anchos=[6.0, 2.6, 2.8, 3.0], size=9, num_cols=(1, 2, 3))

h2("7.3 Apertura por subcategoria")
tabla(["MARCA_ACADEMICA", "MARCA_ACADEMICA_DETALLE", "Clientes", "Obligaciones"],
      [[det_marca[d], d, fmt(det_cli[d]), fmt(det_obl[d])]
       for d, _ in sorted(det_cli.items(), key=lambda x: (det_marca[x[0]], -x[1]))],
      anchos=[5.4, 5.4, 2.2, 2.4], size=8.5, num_cols=(2, 3))

# ================================================================== ANEXO (horizontal)
sec = doc.add_section(WD_SECTION.NEW_PAGE)
sec.orientation = WD_ORIENT.LANDSCAPE
sec.page_width, sec.page_height = sec.page_height, sec.page_width
sec.top_margin = Cm(1.4); sec.bottom_margin = Cm(1.4)
sec.left_margin = Cm(1.4); sec.right_margin = Cm(1.4)

h1(f"Anexo. Las {len(MATRIZ)} casuisticas observadas en la cartera")
par("Salida completa de la consulta marca_academica_combinaciones.sql. Cada fila es una combinacion "
    "real de variables presente en la cartera, con la marca que le asigna la logica propuesta. "
    "Ordenada por marca y por numero de clientes.", size=9.5, color=GRIS)

filas_anexo = []
for c in sorted(MATRIZ, key=lambda x: (x["MARCA_ACADEMICA"], -int(x["CLIENTES_UNICOS"]))):
    filas_anexo.append([
        c["TIPO_CLIENTE"], c["ESTADO_PERIODO"], c["ESTADO_ALUMNO"],
        c["RES_PERFIL_RIESGO"], c["COBERTURA_PERFIL"], c["acceso plataforma"],
        c["ESTADO NOTAS"], c["MARCA_ACADEMICA"], c["MARCA_ACADEMICA_DETALLE"],
        fmt(c["CLIENTES_UNICOS"]), fmt(c["OBLIGACIONES"]),
    ])
tabla(["SEGMENTO", "ESTADO PERIODO", "VINCULO", "RIESGO", "TIENE PERFIL", "ACCESO",
       "NOTAS", "MARCA_ACADEMICA", "MARCA_ACADEMICA_DETALLE", "Clientes", "Oblig."],
      filas_anexo,
      anchos=[2.2, 2.8, 1.9, 1.3, 1.5, 1.3, 1.5, 4.4, 4.6, 1.5, 1.5],
      size=7, num_cols=(9, 10))

doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Analitica financiera - Universidad CUN")
r.font.size = Pt(8.5); r.font.color.rgb = GRIS

try:
    doc.save(SALIDA)
except PermissionError:
    # El .docx esta abierto en Word: se guarda con sufijo para no perder la generacion.
    SALIDA = SALIDA.replace(".docx", "_V2.docx")
    doc.save(SALIDA)
    print("AVISO: el archivo original estaba abierto en Word; se guardo como _V2.")
print("OK ->", SALIDA)
print("Clientes unicos por marca (exclusivo):", {k: fmt(v) for k, v in uni_marca.items()},
      "| suma:", fmt(sum(uni_marca.values())))
